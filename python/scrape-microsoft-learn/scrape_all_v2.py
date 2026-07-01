#!/usr/bin/env python3
"""Microsoft Learn Scraper v2 — downloads learning paths as Markdown, DOCX, TXT, or PDF.

Author:  Matthew Gonzalez Nieves
Version: 1.0.0
Date:    31-03-2025
License: MIT
Website: https://mgntechvault.com

Purpose:
    This tool is intended for personal offline study using AI tools such as
    Microsoft Copilot Notebook and Google NotebookLM. Export your Microsoft Learn
    learning paths and feed them directly to your AI notebook of choice.

    This script is NOT intended to stress or flood Microsoft Learn (DDoS).
    Use it responsibly and only for your own study materials.

Usage:
    python scrape_all_v2.py --format markdown       # required flag
    python scrape_all_v2.py --format docx
    python scrape_all_v2.py --format txt
    python scrape_all_v2.py --format pdf
    python scrape_all_v2.py                         # interactive choice if --format is omitted
    python scrape_all_v2.py --pick                  # interactive learning-path picker
    python scrape_all_v2.py --courses               # interactive course picker (e.g. MD-102)
    python scrape_all_v2.py --locale nl-nl          # picker catalog locale (default en-us)
    python scrape_all_v2.py --config custom.yaml    # use a different config file
    python scrape_all_v2.py --urls URL1 URL2        # override URLs via CLI
    python scrape_all_v2.py --output filename       # override output name (extension auto-determined)
    python scrape_all_v2.py --headed                # show browser window
    python scrape_all_v2.py -v                      # verbose logging

    Without --urls (and without --pick/--courses), an interactive menu asks
    whether to pick from the live Microsoft Learn catalog — individual learning
    paths or a whole course (which bundles several paths) — or use the paths
    from config.yaml.

Requirements:
    pip install playwright pyyaml python-docx
    playwright install chromium
"""

import argparse
import json
import logging
import re
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin

import yaml
from playwright.sync_api import sync_playwright, Page, TimeoutError as PwTimeout

logger = logging.getLogger(__name__)

VALID_FORMATS = ("markdown", "docx", "txt", "pdf")
FORMAT_EXTENSIONS = {"markdown": ".md", "docx": ".docx", "txt": ".txt", "pdf": ".pdf"}

# Browser user-agent reused for scraping and the catalog fetch.
USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)

# Microsoft Learn public catalog API — lists learning paths, courses, modules…
CATALOG_API = "https://learn.microsoft.com/api/catalog/?type={types}&locale={locale}"
DEFAULT_LOCALE = "en-us"
PICKER_RESULT_LIMIT = 40
# Maps a catalog study_guide entry type to the scrape kind used internally.
STUDY_GUIDE_KINDS = {"learningPath": "path", "module": "module"}

# --- Selectors (centralised so changes only need to happen in one place) ---
CONTENT_SELECTOR = "#unit-inner-section"
FALLBACK_CONTENT_SELECTOR = "div[role='main']"
UNIT_LIST_SELECTOR = "#unit-list"
NOISE_SELECTORS = (
    "button", "footer", "nav",
    ".display-flex",
    "#next-section", "#previous-section",
    ".assessment-table",
    ".unit-page-learning-objectives",
)
MIN_CONTENT_LENGTH = 50


# ── Data models ───────────────────────────────────────────────


@dataclass
class Unit:
    title: str
    url: str
    content: str


@dataclass
class Module:
    title: str
    url: str
    units: list[Unit] = field(default_factory=list)


@dataclass
class LearningPath:
    title: str
    url: str
    modules: list[Module] = field(default_factory=list)


@dataclass
class ScrapeStats:
    paths: int = 0
    modules: int = 0
    units_scraped: int = 0
    units_skipped: int = 0
    nav_failures: int = 0
    failed_urls: list[str] = field(default_factory=list)

    def summary(self) -> str:
        lines = [
            f"Paths processed:   {self.paths}",
            f"Modules processed: {self.modules}",
            f"Units scraped:     {self.units_scraped}",
            f"Units skipped:     {self.units_skipped}",
            f"Navigation errors: {self.nav_failures}",
        ]
        if self.failed_urls:
            lines.append("Failed URLs:")
            for url in self.failed_urls:
                lines.append(f"  - {url}")
        return "\n".join(lines)


# ── Utility functions ─────────────────────────────────────────


def clean_text(text: str) -> str:
    """Remove redundant blank lines."""
    if not text:
        return ""
    return re.sub(r"\n\s*\n", "\n\n", text.strip())


# ── Browser helpers ───────────────────────────────────────────


def safe_goto(page: Page, url: str, *, timeout: int = 20_000, retries: int = 3) -> bool:
    """Navigate to a URL with exponential backoff on failure."""
    for attempt in range(retries):
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=timeout)
            return True
        except Exception as e:
            wait = 2 ** attempt
            logger.warning(
                "Navigation failed (%d/%d) %s — %s, retrying in %ds",
                attempt + 1, retries, url, e, wait,
            )
            time.sleep(wait)
    logger.error("Permanently failed: %s", url)
    return False


def wait_for_content(page: Page, timeout: int = 10_000) -> bool:
    """Wait until the content selector is visible on the page."""
    try:
        page.wait_for_selector(CONTENT_SELECTOR, state="attached", timeout=timeout)
        return True
    except PwTimeout:
        logger.debug("Content selector not found within timeout on %s", page.url)
        return False


def dismiss_cookie_banner(page: Page) -> None:
    """Click away the cookie banner if present."""
    for selector in ("button#onetrust-accept-btn-handler", "button.accept-cookies"):
        try:
            page.click(selector, timeout=2000)
            logger.debug("Cookie banner dismissed via %s", selector)
            return
        except PwTimeout:
            continue


def extract_content(page: Page) -> str | None:
    """Extract content via a DOM clone — does not mutate the original."""
    noise_css = ", ".join(NOISE_SELECTORS)
    result = page.evaluate(
        """([contentSel, fallbackSel, noiseSel]) => {
            const el = document.querySelector(contentSel)
                    || document.querySelector(fallbackSel);
            if (!el) return null;
            const clone = el.cloneNode(true);
            clone.querySelectorAll(noiseSel).forEach(e => e.remove());
            return clone.innerText.trim();
        }""",
        [CONTENT_SELECTOR, FALLBACK_CONTENT_SELECTOR, noise_css],
    )
    if result and len(result) > MIN_CONTENT_LENGTH:
        return clean_text(result)
    return None


# ── URL extraction ────────────────────────────────────────────


def normalize_url(base: str, href: str) -> str:
    """Build a full URL without query params, with trailing slash stripped."""
    return urljoin(base, href).split("?")[0].rstrip("/")


def get_module_urls(page: Page) -> list[str]:
    """Retrieve unique module URLs from a learning path page."""
    links = page.query_selector_all("a[href*='/modules/']")
    seen: set[str] = set()
    urls: list[str] = []
    for link in links:
        href = link.get_attribute("href")
        if not href:
            continue
        full = normalize_url(page.url, href)
        if full not in seen:
            seen.add(full)
            urls.append(full)
    return urls


def get_unit_urls(page: Page, module_url: str) -> list[str]:
    """Retrieve unique unit URLs from a module page."""
    container = page.query_selector(UNIT_LIST_SELECTOR)
    if not container:
        logger.debug("Selector '%s' not found, falling back to <main>", UNIT_LIST_SELECTOR)
        container = page.query_selector("main")
    if not container:
        return []

    normalized_module = module_url.rstrip("/")
    seen: set[str] = set()
    urls: list[str] = []
    for link in container.query_selector_all("a"):
        href = link.get_attribute("href")
        if not href:
            continue
        full = normalize_url(page.url, href)
        # Skip the module link itself and external links
        if full == normalized_module:
            continue
        if full not in seen:
            seen.add(full)
            urls.append(full)
    return urls


# ── Core scraping ─────────────────────────────────────────────


def scrape_unit(page: Page, url: str) -> Unit | None:
    """Scrape a single unit (lesson) and return a Unit, or None on failure."""
    if not safe_goto(page, url):
        return None
    wait_for_content(page)

    h1 = page.query_selector("h1")
    title = h1.inner_text().strip() if h1 else "Unknown lesson"
    content = extract_content(page)

    if not content:
        logger.info("  No usable content: %s", url)
        return None
    return Unit(title=title, url=url, content=content)


def scrape_module(page: Page, url: str, stats: ScrapeStats, seen_urls: set[str]) -> Module | None:
    """Scrape a module including all its units."""
    if not safe_goto(page, url):
        stats.nav_failures += 1
        stats.failed_urls.append(url)
        return None

    h1 = page.query_selector("h1")
    title = h1.inner_text().strip() if h1 else "Unknown module"
    module = Module(title=title, url=url)

    unit_urls = get_unit_urls(page, url)
    logger.info("    %d units found in '%s'", len(unit_urls), title)

    for idx, u_url in enumerate(unit_urls, 1):
        if u_url in seen_urls:
            logger.debug("    Unit already seen, skipping: %s", u_url)
            continue
        seen_urls.add(u_url)

        logger.info("      [%d/%d] %s", idx, len(unit_urls), u_url.split("/")[-1])
        unit = scrape_unit(page, u_url)
        if unit:
            module.units.append(unit)
            stats.units_scraped += 1
        else:
            stats.units_skipped += 1

    stats.modules += 1
    return module


def scrape_path(page: Page, url: str, stats: ScrapeStats, seen_urls: set[str]) -> LearningPath | None:
    """Scrape a learning path including all its modules and units."""
    logger.info("Learning path: %s", url.split("/")[-2])
    if not safe_goto(page, url):
        stats.nav_failures += 1
        stats.failed_urls.append(url)
        return None

    h1 = page.query_selector("h1")
    title = h1.inner_text().strip() if h1 else "Unknown path"
    path = LearningPath(title=title, url=url)

    module_urls = get_module_urls(page)
    logger.info("  %d modules found in '%s'", len(module_urls), title)

    for m_idx, m_url in enumerate(module_urls, 1):
        logger.info("  [%d/%d] Processing module...", m_idx, len(module_urls))
        module = scrape_module(page, m_url, stats, seen_urls)
        if module and module.units:
            path.modules.append(module)

    stats.paths += 1
    return path


# ── Writers ───────────────────────────────────────────────────


def _atomic_write_text(tmp: Path, output: Path, content_fn) -> None:
    """Write via a temporary file and atomically rename on success."""
    with open(tmp, "w", encoding="utf-8") as f:
        content_fn(f)
    tmp.replace(output)


def write_markdown(
    paths: list[LearningPath],
    output: Path,
    title: str = "AZ-305 Course Material",
) -> None:
    """Write to Markdown (.md). Atomic write via .tmp file."""
    tmp = output.with_suffix(".tmp")

    def _write(f):
        f.write(f"# {title}\n\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")

        f.write("## Table of Contents\n\n")
        for lp in paths:
            f.write(f"- **{lp.title}**\n")
            for mod in lp.modules:
                f.write(f"  - {mod.title} ({len(mod.units)} lessons)\n")
        f.write("\n---\n\n")

        for lp in paths:
            f.write(f"## {lp.title}\n\n")
            f.write(f"Source: {lp.url}\n\n")
            for mod in lp.modules:
                f.write(f"### {mod.title}\n\n")
                for unit in mod.units:
                    f.write(f"#### {unit.title}\n\n")
                    f.write(f"{unit.content}\n\n")
                    f.write("---\n\n")
            f.flush()

    _atomic_write_text(tmp, output, _write)
    logger.info("Markdown written to %s", output)


def write_txt(
    paths: list[LearningPath],
    output: Path,
    title: str = "AZ-305 Course Material",
) -> None:
    """Write to plain text (.txt). Atomic write via .tmp file."""
    tmp = output.with_suffix(".tmp")
    separator = "=" * 72
    thin_sep = "-" * 72

    def _write(f):
        f.write(f"{title}\n")
        f.write(f"{separator}\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")

        f.write("TABLE OF CONTENTS\n")
        f.write(f"{thin_sep}\n")
        for lp in paths:
            f.write(f"  {lp.title}\n")
            for mod in lp.modules:
                f.write(f"    - {mod.title} ({len(mod.units)} lessons)\n")
        f.write(f"\n{separator}\n\n")

        for lp in paths:
            f.write(f"{lp.title.upper()}\n")
            f.write(f"{separator}\n")
            f.write(f"Source: {lp.url}\n\n")
            for mod in lp.modules:
                f.write(f"{mod.title}\n")
                f.write(f"{thin_sep}\n\n")
                for unit in mod.units:
                    f.write(f"  {unit.title}\n")
                    f.write(f"  {'·' * (len(unit.title) + 2)}\n\n")
                    # Indent each content line
                    for line in unit.content.splitlines():
                        f.write(f"  {line}\n")
                    f.write(f"\n{thin_sep}\n\n")
            f.flush()

    _atomic_write_text(tmp, output, _write)
    logger.info("TXT written to %s", output)


def write_docx(
    paths: list[LearningPath],
    output: Path,
    title: str = "AZ-305 Course Material",
) -> None:
    """Write to a Word document (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error(
            "python-docx is not installed. Run: pip install python-docx"
        )
        raise SystemExit(1)

    doc = Document()

    # Document title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Generation date
    date_para = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Table of contents (text-based)
    doc.add_heading("Table of Contents", level=1)
    for lp in paths:
        toc_path = doc.add_paragraph(style="List Bullet")
        toc_path.add_run(lp.title).bold = True
        for mod in lp.modules:
            toc_mod = doc.add_paragraph(style="List Bullet 2")
            toc_mod.add_run(f"{mod.title} ({len(mod.units)} lessons)")

    doc.add_page_break()

    # Content
    for lp in paths:
        doc.add_heading(lp.title, level=1)
        source_para = doc.add_paragraph()
        source_para.add_run("Source: ").bold = True
        source_para.add_run(lp.url)

        for mod in lp.modules:
            doc.add_heading(mod.title, level=2)

            for unit in mod.units:
                doc.add_heading(unit.title, level=3)
                # Split content on blank lines — each block becomes its own paragraph
                for block in unit.content.split("\n\n"):
                    block = block.strip()
                    if block:
                        doc.add_paragraph(block)

                doc.add_paragraph("─" * 40)

    tmp = output.with_suffix(".tmp")
    doc.save(tmp)
    tmp.replace(output)
    logger.info("DOCX written to %s", output)


PDF_STYLE = """
    body { font-family: -apple-system, Segoe UI, Helvetica, Arial, sans-serif;
           font-size: 11pt; line-height: 1.5; color: #1b1b1b; margin: 2em; }
    h1 { font-size: 22pt; border-bottom: 2px solid #0067b8; padding-bottom: .2em; }
    h2 { font-size: 17pt; color: #0067b8; margin-top: 1.4em;
         page-break-before: always; }
    h2:first-of-type { page-break-before: avoid; }
    h3 { font-size: 14pt; margin-top: 1em; }
    h4 { font-size: 12pt; margin-top: .8em; color: #333; }
    .meta { color: #666; font-size: 9pt; }
    .source { color: #666; font-size: 9pt; word-break: break-all; }
    .toc a, .content { color: inherit; }
    pre { white-space: pre-wrap; font-family: inherit; }
    hr { border: none; border-top: 1px solid #ddd; margin: 1em 0; }
"""


def _html_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _render_html(paths: list[LearningPath], title: str) -> str:
    """Build a single self-contained HTML document for PDF rendering."""
    out: list[str] = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>",
        f"<title>{_html_escape(title)}</title>",
        f"<style>{PDF_STYLE}</style></head><body>",
        f"<h1>{_html_escape(title)}</h1>",
        f"<p class='meta'>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>",
        "<h2 style='page-break-before:avoid'>Table of Contents</h2><ul class='toc'>",
    ]
    for lp in paths:
        out.append(f"<li><strong>{_html_escape(lp.title)}</strong><ul>")
        for mod in lp.modules:
            out.append(
                f"<li>{_html_escape(mod.title)} ({len(mod.units)} lessons)</li>"
            )
        out.append("</ul></li>")
    out.append("</ul>")

    for lp in paths:
        out.append(f"<h2>{_html_escape(lp.title)}</h2>")
        out.append(f"<p class='source'>Source: {_html_escape(lp.url)}</p>")
        for mod in lp.modules:
            out.append(f"<h3>{_html_escape(mod.title)}</h3>")
            for unit in mod.units:
                out.append(f"<h4>{_html_escape(unit.title)}</h4>")
                out.append(f"<pre class='content'>{_html_escape(unit.content)}</pre>")
                out.append("<hr>")
    out.append("</body></html>")
    return "".join(out)


def write_pdf(
    paths: list[LearningPath],
    output: Path,
    title: str = "AZ-305 Course Material",
) -> None:
    """Write to PDF (.pdf) by rendering HTML in headless Chromium."""
    html = _render_html(paths, title)
    tmp = output.with_suffix(".tmp")
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_content(html, wait_until="load")
        page.pdf(
            path=str(tmp),
            format="A4",
            print_background=True,
            margin={"top": "1.5cm", "bottom": "1.5cm", "left": "1.2cm", "right": "1.2cm"},
        )
        browser.close()
    tmp.replace(output)
    logger.info("PDF written to %s", output)


def write_output(
    paths: list[LearningPath],
    output: Path,
    fmt: str,
    title: str,
) -> None:
    """Dispatcher: select the correct writer based on the specified format."""
    writers = {
        "markdown": write_markdown,
        "txt": write_txt,
        "docx": write_docx,
        "pdf": write_pdf,
    }
    writers[fmt](paths, output, title)


# ── Config & CLI ──────────────────────────────────────────────


def load_config(path: str) -> dict:
    """Load YAML configuration. Returns empty dict if file does not exist."""
    try:
        with open(path) as f:
            return yaml.safe_load(f) or {}
    except FileNotFoundError:
        logger.debug("Config file '%s' not found, using defaults", path)
        return {}


# ── Catalog fetch & picker ────────────────────────────────────


def _as_path_targets(urls: list[str]) -> list[dict]:
    """Wrap plain learning-path URLs as scrape targets."""
    return [{"url": u, "kind": "path"} for u in urls]


def _slugify(text: str) -> str:
    """Turn arbitrary text into a filesystem-safe basename."""
    slug = re.sub(r"[^\w\-]+", "-", (text or "").strip()).strip("-")
    return slug or "output"


def _fetch_catalog_raw(types: str, locale: str) -> dict:
    """Fetch raw catalog JSON for the given comma-separated catalog types.

    Uses a real headless Chromium page because a plain HTTP request is blocked
    (HTTP 403) by Microsoft Learn's bot protection.
    """
    url = CATALOG_API.format(types=types, locale=locale)
    logger.info("Fetching catalog (%s, %s)…", types, locale)
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(user_agent=USER_AGENT)
        page = context.new_page()
        try:
            if not safe_goto(page, url):
                return {}
            raw = page.evaluate("() => document.body.innerText")
        finally:
            browser.close()

    try:
        return json.loads(raw)
    except (json.JSONDecodeError, TypeError):
        logger.error("Could not parse catalog response as JSON.")
        return {}


def fetch_learning_paths(locale: str = DEFAULT_LOCALE) -> list[dict]:
    """Fetch every learning path from the Microsoft Learn catalog API."""
    data = _fetch_catalog_raw("learningPaths", locale)
    paths = data.get("learningPaths", [])
    paths.sort(key=lambda lp: (lp.get("title") or "").lower())
    logger.info("Catalog loaded: %d learning paths available.", len(paths))
    return paths


def fetch_courses(locale: str = DEFAULT_LOCALE) -> list[dict]:
    """Fetch every course plus the learning paths/modules they reference.

    A course has no lessons of its own — its ``study_guide`` references learning
    paths and modules by uid. We fetch those alongside the courses in one call
    and attach a resolved ``targets`` list ([{url, kind}]) to each course so the
    scraper knows exactly what to download.
    """
    data = _fetch_catalog_raw("courses,learningPaths,modules", locale)
    courses = data.get("courses", [])
    uid_to_url = {
        item["uid"]: item.get("url")
        for item in data.get("learningPaths", []) + data.get("modules", [])
        if item.get("uid")
    }
    for course in courses:
        targets = []
        for entry in course.get("study_guide") or []:
            url = uid_to_url.get(entry.get("uid"))
            kind = STUDY_GUIDE_KINDS.get(entry.get("type"))
            if url and kind:
                targets.append({"url": url, "kind": kind})
        course["targets"] = targets
    courses.sort(key=lambda c: (c.get("title") or "").lower())
    logger.info("Catalog loaded: %d courses available.", len(courses))
    return courses


def _select_from_catalog(
    items: list[dict],
    *,
    noun: str,
    label_prefix=None,
    label_extra=None,
) -> list[dict]:
    """Shared search + multi-select UI over catalog items.

    ``noun`` names the item type (e.g. 'learning path'). ``label_prefix`` is an
    optional callable item -> str shown before each title (e.g. a course
    number); ``label_extra`` is appended after each title (e.g. scope hints).
    Returns the selected item dicts, or [] if the user cancels.
    """

    def _label(it: dict) -> str:
        prefix = label_prefix(it) if label_prefix else ""
        prefix = f"{prefix} — " if prefix else ""
        extra = f"  ({label_extra(it)})" if label_extra else ""
        return f"{prefix}{it.get('title', '(untitled)')}{extra}"

    if not items:
        print(f"\nNo {noun}s could be retrieved from the catalog.")
        return []

    print()
    print(f"{noun.capitalize()} picker")
    print("=" * 42)
    print(f"{len(items)} {noun}s available.")

    def _ask(prompt: str) -> str:
        try:
            return input(prompt).strip()
        except (EOFError, KeyboardInterrupt):
            print("\nCancelled.")
            raise SystemExit(0)

    while True:
        term = _ask(
            f"\nSearch by keyword (e.g. 'azure'), or Enter to list all {noun}s: "
        ).lower()

        if term:
            matches = [
                it for it in items
                if term in (it.get("title") or "").lower()
                or term in (it.get("summary") or "").lower()
            ]
        else:
            matches = items

        if not matches:
            print(f"  No {noun}s match '{term}'. Try another keyword.")
            continue

        if len(matches) > PICKER_RESULT_LIMIT:
            print(
                f"  {len(matches)} matches — showing first {PICKER_RESULT_LIMIT}. "
                "Refine your keyword to narrow the list."
            )
            shown = matches[:PICKER_RESULT_LIMIT]
        else:
            shown = matches

        print()
        for idx, it in enumerate(shown, 1):
            print(f"  {idx:>3}  {_label(it)}")
        print()

        choice = _ask(
            "Enter number(s) to select (comma-separated), "
            "or 's' to search again, 'q' to quit: "
        ).lower()

        if choice in ("q", "quit"):
            print("Cancelled.")
            return []
        if choice in ("", "s", "search"):
            continue

        try:
            indices = [int(c) for c in choice.replace(" ", "").split(",") if c]
        except ValueError:
            print("  Invalid input — enter numbers like '1' or '1,3,5'.")
            continue

        selected = [shown[i - 1] for i in indices if 1 <= i <= len(shown)]
        if not selected:
            print("  No valid numbers in range — try again.")
            continue

        print("\n-> Selected:")
        for it in selected:
            print(f"     • {_label(it)}")
        print()
        return selected


def pick_learning_paths(
    paths: list[dict],
) -> tuple[list[dict], str | None, str | None]:
    """Search & select learning paths.

    Returns (targets, default title, suggested output basename). The basename is
    None so the existing config/output naming is left untouched.
    """
    selected = _select_from_catalog(paths, noun="learning path")
    targets = _as_path_targets([lp["url"] for lp in selected if lp.get("url")])
    title = selected[0].get("title") if selected else None
    return targets, title, None


def pick_courses(
    courses: list[dict],
) -> tuple[list[dict], str | None, str | None]:
    """Search & select courses. Expands each course's study_guide into targets.

    Returns (targets, default title, suggested output basename). For a course
    the title and filename are derived from its course number (e.g. MD-102T00),
    falling back to a slug of the title when the number is missing. A single
    course can expand to many learning paths/modules, so the scrape may be large.
    """
    selected = _select_from_catalog(
        courses,
        noun="course",
        label_prefix=lambda c: c.get("course_number") or "",
        label_extra=lambda c: f"{len(c.get('targets', []))} paths/modules",
    )
    targets: list[dict] = []
    for course in selected:
        targets.extend(course.get("targets", []))
    if selected and not targets:
        print("  Selected course(s) have no scrapeable content.")

    if not selected:
        return targets, None, None

    first = selected[0]
    number = (first.get("course_number") or "").strip()
    raw_title = first.get("title") or "Course"
    title = f"{number} — {raw_title}" if number else raw_title
    basename = _slugify(number) if number else _slugify(raw_title)
    return targets, title, basename


def ask_source_interactively(config_paths: list[str], config_name: str) -> str:
    """Ask where the learning content should come from.

    Returns 'paths' (live learning paths), 'courses' (live courses), or
    'config' (config YAML). The config option is only offered when the config
    file actually contains paths.
    """
    has_config = bool(config_paths)
    print()
    print("Where should the content come from?")
    print("=" * 42)
    print("  1  ->  Live learning paths  (search & pick individual paths)")
    print("  2  ->  Live courses         (exam-style, bundles paths, e.g. MD-102)")
    if has_config:
        print(f"  3  ->  Config file          ({config_name}, {len(config_paths)} path(s))")
    else:
        print(f"  3  ->  Config file          ({config_name} — no paths found)")
    print()

    while True:
        try:
            choice = input(
                "Enter 1, 2 or 3 (or 'paths' / 'courses' / 'config'): "
            ).strip().lower()
        except (EOFError, KeyboardInterrupt):
            print("\nCancelled.")
            raise SystemExit(0)

        if choice in ("1", "paths", "path", "learningpaths"):
            print("-> Source: live learning paths\n")
            return "paths"
        if choice in ("2", "courses", "course"):
            print("-> Source: live courses\n")
            return "courses"
        if choice in ("3", "config", "yaml"):
            if not has_config:
                print(
                    f"  No paths found in {config_name}. "
                    "Choose 1 or 2 to use the live catalog."
                )
                continue
            print(f"-> Source: config file ({config_name})\n")
            return "config"

        print("  Invalid choice — enter 1 (paths), 2 (courses) or 3 (config).")


def ask_format_interactively() -> str:
    """Ask the user interactively for an output format when --format is not provided."""
    options = {"1": "markdown", "2": "docx", "3": "txt", "4": "pdf"}
    print()
    print("Choose an output format")
    print("=" * 42)
    print("  1  ->  Markdown  (.md)")
    print("  2  ->  Word      (.docx)")
    print("  3  ->  Plain text (.txt)")
    print("  4  ->  PDF       (.pdf)")
    print()

    while True:
        try:
            choice = input(
                "Enter 1, 2, 3 or 4 (or type 'markdown', 'docx', 'txt', 'pdf'): "
            ).strip().lower()
        except (EOFError, KeyboardInterrupt):
            print("\nCancelled.")
            raise SystemExit(0)

        # Accept both number and name
        if choice in options:
            selected = options[choice]
            print(f"-> Selected format: {selected}\n")
            return selected
        if choice in VALID_FORMATS:
            print(f"-> Selected format: {choice}\n")
            return choice

        print(f"  Invalid choice '{choice}'. Choose 1, 2, 3, 4 or type the format name.")


def resolve_format(raw: str | None) -> str:
    """Validate the format. If None: ask interactively. If invalid: error + exit."""
    if raw is None:
        return ask_format_interactively()
    fmt = raw.strip().lower()
    if fmt not in VALID_FORMATS:
        print(
            f"\nError: '{raw}' is not a valid format.\n"
            f"Choose one of: {', '.join(VALID_FORMATS)}\n\n"
            f"Usage:  python scrape_all_v2.py --format markdown\n"
            f"        python scrape_all_v2.py --format docx\n"
            f"        python scrape_all_v2.py --format txt\n"
            f"        python scrape_all_v2.py --format pdf\n",
            file=sys.stderr,
        )
        raise SystemExit(1)
    return fmt


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description=(
            "Microsoft Learn Scraper — downloads learning paths as Markdown, DOCX, TXT, or PDF.\n\n"
            "  --format is required. If omitted you will be prompted interactively.\n"
            "  When no URLs are given, a menu asks: live learning paths, live\n"
            "  courses (e.g. MD-102), or the paths from config.yaml."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python scrape_all_v2.py --format markdown            # source menu, then pick\n"
            "  python scrape_all_v2.py --format pdf --pick           # force learning-path picker\n"
            "  python scrape_all_v2.py --format pdf --courses        # force course picker\n"
            "  python scrape_all_v2.py --format docx --output MyCourse --urls URL\n"
            "  python scrape_all_v2.py --format txt --config other.yaml -v\n"
        ),
    )
    p.add_argument(
        "--format",
        dest="fmt",
        choices=VALID_FORMATS,
        metavar="FORMAT",
        help=f"Output format — required. Choose from: {', '.join(VALID_FORMATS)}",
    )
    p.add_argument("--config", default="config.yaml", help="YAML config file (default: config.yaml)")
    p.add_argument("--urls", nargs="+", help="Learning path URLs (overrides config)")
    p.add_argument(
        "--pick",
        action="store_true",
        help="Skip the source menu and go straight to the live learning-path picker",
    )
    p.add_argument(
        "--courses",
        action="store_true",
        help="Skip the source menu and go straight to the live course picker "
             "(a course bundles several learning paths, e.g. MD-102)",
    )
    p.add_argument(
        "--locale",
        default=DEFAULT_LOCALE,
        help=f"Catalog locale for the picker (default: {DEFAULT_LOCALE}, e.g. nl-nl)",
    )
    p.add_argument("--output", help="Output filename without extension, or full path (extension is auto-determined)")
    p.add_argument("--title", help="Title of the output document")
    p.add_argument("--headed", action="store_true", help="Show the browser window (for debugging)")
    p.add_argument("-v", "--verbose", action="store_true", help="Verbose logging (debug level)")
    return p.parse_args()


# ── Main ──────────────────────────────────────────────────────


def main() -> None:
    args = parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler("scraper.log", encoding="utf-8"),
        ],
    )

    # ── Determine format (required — interactive if not provided) ─────────
    fmt = resolve_format(args.fmt)

    # ── Load config — CLI args take priority over YAML ────────────────────
    cfg = load_config(args.config)
    config_paths = cfg.get("paths", [])
    title = args.title or cfg.get("title", "AZ-305 Course Material")

    # ── Decide where the content comes from ───────────────────────────────
    #   --urls    : explicit CLI URLs always win, no prompt
    #   --pick    : force the live learning-path picker
    #   --courses : force the live course picker
    #   otherwise : show a TUI menu (paths / courses / config YAML).
    #               When stdin is not a TTY (cron/pipe) we can't prompt, so we
    #               fall back to the config paths to preserve automation.
    if args.urls:
        source = "urls"
    elif args.pick:
        source = "paths"
    elif args.courses:
        source = "courses"
    elif sys.stdin.isatty():
        source = ask_source_interactively(config_paths, args.config)
    else:
        logger.info("Non-interactive session — using paths from %s.", args.config)
        source = "config"

    picked_title: str | None = None
    picked_output: str | None = None
    if source == "urls":
        targets = _as_path_targets(args.urls)
    elif source == "paths":
        targets, picked_title, picked_output = pick_learning_paths(
            fetch_learning_paths(args.locale)
        )
    elif source == "courses":
        targets, picked_title, picked_output = pick_courses(
            fetch_courses(args.locale)
        )
    else:  # config
        targets = _as_path_targets(config_paths)

    if not targets:
        if source == "config":
            logger.error("No paths found in %s — nothing to do.", args.config)
            raise SystemExit(1)
        logger.info("Nothing selected — nothing to do.")
        raise SystemExit(0)

    # ── Document title ────────────────────────────────────────────────────
    #   --title wins; then a picked course title (incl. its number) overrides
    #   the config title; a picked path title only fills in when config has none.
    if args.title:
        pass
    elif source == "courses" and picked_title:
        title = picked_title
    elif picked_title and not cfg.get("title"):
        title = picked_title

    # ── Output filename — extension always follows the format ─────────────
    #   --output wins; then a course-derived name (e.g. MD-102T00); then config.
    raw_output = args.output or picked_output or cfg.get("output", "output")
    output_path = Path(raw_output).with_suffix(FORMAT_EXTENSIONS[fmt])
    if len(output_path.parts) == 1:
        output_path = Path("exports") / output_path
    output_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info(
        "Starting scraper — format: %s | %d target(s) | output: %s",
        fmt, len(targets), output_path,
    )

    stats = ScrapeStats()
    seen_urls: set[str] = set()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=not args.headed)
        context = browser.new_context(user_agent=USER_AGENT)
        page = context.new_page()

        # Handle cookie banner on first navigation
        safe_goto(page, targets[0]["url"])
        dismiss_cookie_banner(page)

        results: list[LearningPath] = []
        for target in targets:
            if target["kind"] == "module":
                # A standalone module (from a course): wrap it in a one-module
                # learning path so the writers keep their path → module → unit shape.
                module = scrape_module(page, target["url"], stats, seen_urls)
                if module and module.units:
                    results.append(
                        LearningPath(title=module.title, url=target["url"], modules=[module])
                    )
            else:
                lp = scrape_path(page, target["url"], stats, seen_urls)
                if lp and lp.modules:
                    results.append(lp)

        browser.close()

    if results:
        write_output(results, output_path, fmt=fmt, title=title)
    else:
        logger.warning("No results to write.")

    # Summary
    logger.info("Done!\n%s", stats.summary())


if __name__ == "__main__":
    main()
